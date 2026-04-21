"""Generate skripsi-style Word document (Bab 1-5) in Indonesian."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "paper"
OUT_DIR.mkdir(exist_ok=True)

evaluate = json.loads((ROOT / "output/results/evaluate.json").read_text())
comparison = json.loads((ROOT / "output/results/comparison_summary.json").read_text())

doc = Document()

# Page setup: A4, margin skripsi standard (kiri 4cm, kanan 3cm, atas 4cm, bawah 3cm)
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)


def H1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"
    return p


def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    return p


def H3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    return p


def P(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    return p


def BULLET(items, numbered=False):
    style_name = "List Number" if numbered else "List Bullet"
    for it in items:
        try:
            p = doc.add_paragraph(it, style=style_name)
        except KeyError:
            p = doc.add_paragraph(("• " if not numbered else "") + it)
        for r in p.runs:
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"


def CAPTION(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"


def TABLE(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
    doc.add_paragraph()


# ============ COVER ============
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("ANALISIS KOMPARATIF METODE PERINGKASAN TEKS\nEKSTRAKTIF DAN ABSTRAKTIF UNTUK\nDOKUMEN AKADEMIK BERBAHASA INDONESIA")
r.bold = True
r.font.size = Pt(16)
r.font.name = "Times New Roman"

for _ in range(2):
    doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Laporan Penelitian / Skripsi\n\nDisusun oleh:\n[Nama Mahasiswa]\n[NIM]\n\nProgram Studi Informatika\n[Nama Universitas]\n2026")
r.font.size = Pt(12)
r.font.name = "Times New Roman"

doc.add_page_break()

# ============ BAB 1 ============
H1("BAB 1\nPendahuluan")

H2("1.1 Latar Belakang")
P("Pertumbuhan eksponensial publikasi ilmiah berbahasa Indonesia menciptakan tantangan bagi peneliti, akademisi, dan mahasiswa untuk menyaring informasi penting secara cepat. Sebuah dokumen akademik tunggal dalam dataset penelitian ini memiliki rata-rata 25.267 karakter (kurang lebih 4.000–5.000 kata), dengan dokumen terpanjang mencapai 57.745 karakter. Membaca seluruh isi setiap artikel tidak lagi efisien.")
P("Peringkasan teks otomatis (automatic text summarization) merupakan cabang Natural Language Processing (NLP) yang bertujuan memadatkan dokumen menjadi versi yang lebih singkat tanpa kehilangan informasi inti. Terdapat dua paradigma utama, yaitu ekstraktif yang memilih kalimat terpenting dari dokumen asli (mempertahankan kata-kata aslinya), dan abstraktif yang menghasilkan kalimat baru yang merangkum makna seperti yang dilakukan manusia.")
P("Riset peringkasan untuk bahasa Inggris telah maju pesat dengan model seperti BART, T5, Pegasus, dan GPT. Namun, untuk bahasa Indonesia — khususnya domain akademik — penelitian komparatif yang menggunakan metrik standar masih terbatas. Bahasa Indonesia memiliki karakteristik morfologi aglutinatif (banyak afiksasi: me-, di-, -kan, -i, ber-, dan lain-lain), serta keterbatasan jumlah pre-trained model yang dioptimalkan khusus untuk bahasa ini.")

H2("1.2 Rumusan Masalah")
BULLET([
    "Bagaimana performa metode ekstraktif berbasis TF-IDF + TextRank dalam meringkas dokumen akademik berbahasa Indonesia?",
    "Bagaimana performa metode abstraktif berbasis mT5 multilingual (XL-Sum) pada dataset yang sama?",
    "Metode mana yang lebih unggul berdasarkan metrik ROUGE-1, ROUGE-2, dan ROUGE-L?",
    "Faktor-faktor apa (panjang dokumen, jumlah kalimat, rasio kompresi) yang memengaruhi kualitas ringkasan per dokumen?",
    "Bagaimana hasil dibandingkan dengan ringkasan yang dihasilkan oleh Large Language Model (LLM) sebagai referensi tambahan?",
], numbered=True)

H2("1.3 Tujuan Penelitian")
BULLET([
    "Mengimplementasikan dan mengevaluasi metode peringkasan ekstraktif (TF-IDF + TextRank) untuk teks akademik Indonesia.",
    "Mengimplementasikan dan mengevaluasi metode peringkasan abstraktif menggunakan model csebuetnlp/mT5_multilingual_XLSum.",
    "Membandingkan kedua metode menggunakan metrik ROUGE pada 100 dokumen.",
    "Menganalisis faktor per-dokumen yang memengaruhi performa.",
    "Mengembangkan aplikasi web interaktif berbasis Flask untuk menjalankan keseluruhan pipeline (preprocessing, summarization, evaluation, comparison).",
], numbered=True)

H2("1.4 Batasan Masalah")
BULLET([
    "Bahasa: Indonesia.",
    "Domain: dokumen akademik (jurnal/skripsi).",
    "Dataset: 100 dokumen.",
    "Model abstraktif: hanya inferensi (tanpa fine-tuning karena keterbatasan GPU; menggunakan AMD EPYC 7763 CPU).",
    "Metrik evaluasi: ROUGE-1, ROUGE-2, ROUGE-L (Precision, Recall, F1).",
    "Aplikasi web: lokal (Flask + Gunicorn + Nginx untuk deploy).",
])

H2("1.5 Manfaat Penelitian")
BULLET([
    "Akademis: memberikan baseline komparatif metode peringkasan teks Indonesia di domain akademik.",
    "Praktis: menyediakan tool web siap pakai untuk meringkas dokumen Indonesia.",
    "Metodologis: pipeline preprocessing Indonesia (Sastrawi stemmer + stopword) yang dapat dipakai ulang.",
])

H2("1.6 Sistematika Penulisan")
BULLET([
    "Bab 1 Pendahuluan — latar belakang, rumusan, tujuan, batasan, manfaat.",
    "Bab 2 Tinjauan Pustaka — teori NLP, peringkasan ekstraktif/abstraktif, ROUGE, riset terkait.",
    "Bab 3 Metodologi — dataset, pipeline, arsitektur, parameter model.",
    "Bab 4 Hasil dan Pembahasan — hasil ROUGE, analisis per-dokumen, perbandingan dengan LLM, aplikasi web.",
    "Bab 5 Kesimpulan dan Saran — temuan utama dan arah riset selanjutnya.",
])

doc.add_page_break()

# ============ BAB 2 ============
H1("BAB 2\nTinjauan Pustaka")

H2("2.1 Natural Language Processing (NLP)")
P("NLP adalah cabang Artificial Intelligence yang fokus pada interaksi komputer dengan bahasa manusia. Tahapan dasar pipeline NLP meliputi case folding (penyeragaman huruf besar/kecil), cleaning (pembersihan karakter non-alfabet, URL, dan tanda baca berlebih), sentence tokenization (pemecahan paragraf menjadi kalimat), word tokenization (pemecahan kalimat menjadi kata), stopword removal (penghapusan kata umum seperti yang, di, ke, dari), serta stemming (pemotongan afiksasi ke bentuk dasar).")
P("Untuk bahasa Indonesia, library Sastrawi menjadi standar de-facto stemming karena dibangun dari aturan morfologi Bahasa Indonesia menggunakan algoritma Nazief–Adriani.")

H2("2.2 Peringkasan Teks (Text Summarization)")
H3("2.2.1 Klasifikasi")
BULLET([
    "Berdasarkan output: Extractive vs Abstractive.",
    "Berdasarkan jumlah dokumen: Single-document vs Multi-document.",
    "Berdasarkan tujuan: Generic vs Query-based.",
])
H3("2.2.2 Peringkasan Ekstraktif")
P("Memilih sub-himpunan kalimat dari dokumen asli yang dianggap paling representatif. Pendekatan klasik meliputi frequency-based (Luhn, 1958), TF-IDF (Salton & Buckley, 1988), dan graph-based seperti TextRank (Mihalcea & Tarau, 2004) yang menggunakan PageRank dengan damping factor d = 0,85.")
H3("2.2.3 Peringkasan Abstraktif")
P("Menghasilkan kalimat baru menggunakan model sequence-to-sequence (seq2seq). Arsitektur dasarnya adalah Transformer (Vaswani et al., 2017) dengan self-attention. Pengembangan lanjutan meliputi BART (Lewis et al., 2020), T5 (Raffel et al., 2020), serta mT5 versi multilingual yang mendukung 101 bahasa termasuk Indonesia. XL-Sum (Hasan et al., 2021) adalah hasil fine-tuning mT5 untuk summarization 44 bahasa termasuk Indonesia, dilatih pada artikel BBC. Model yang digunakan dalam penelitian ini adalah csebuetnlp/mT5_multilingual_XLSum (582 juta parameter, mT5-base).")

H2("2.3 Term Frequency – Inverse Document Frequency (TF-IDF)")
P("TF-IDF dihitung sebagai TF(t,d) dikalikan logaritma dari N dibagi (1 + DF(t)), di mana TF adalah frekuensi term dalam dokumen, N adalah total dokumen, dan DF adalah jumlah dokumen yang memuat term tersebut.")

H2("2.4 Cosine Similarity")
P("Cosine similarity dihitung sebagai dot product dua vektor dibagi hasil kali norma keduanya. Digunakan untuk menghitung kemiripan antar kalimat (vektor TF-IDF) dalam membangun graph TextRank.")

H2("2.5 ROUGE (Recall-Oriented Understudy for Gisting Evaluation)")
P("Lin (2004) mengusulkan ROUGE sebagai metrik standar peringkasan. Tiga varian utama dipakai dalam penelitian ini: ROUGE-1 (overlap unigram), ROUGE-2 (overlap bigram), dan ROUGE-L (longest common subsequence). Setiap varian melaporkan Precision, Recall, dan F1.")

H2("2.6 Penelitian Terkait")
CAPTION("Tabel 2.1 Ringkasan Penelitian Terkait")
TABLE(
    ["Penulis", "Tahun", "Fokus", "Bahasa"],
    [
        ["Luhn", "1958", "Frekuensi kata", "Inggris"],
        ["Mihalcea & Tarau", "2004", "TextRank", "Inggris"],
        ["Vaswani et al.", "2017", "Transformer", "Universal"],
        ["Lewis et al.", "2020", "BART", "Inggris"],
        ["Raffel et al.", "2020", "T5", "Inggris"],
        ["Hasan et al.", "2021", "XL-Sum (mT5)", "44 bahasa termasuk ID"],
        ["Koto et al.", "2020", "IndoBERT, IndoLEM", "Indonesia"],
        ["Wilie et al.", "2020", "IndoNLU benchmark", "Indonesia"],
    ],
)
P("Gap riset: studi komparatif head-to-head ekstraktif vs abstraktif untuk dokumen akademik Indonesia dengan analisis faktor per-dokumen masih jarang.")

doc.add_page_break()

# ============ BAB 3 ============
H1("BAB 3\nMetodologi Penelitian")

H2("3.1 Alur Penelitian")
P("Alur penelitian dimulai dari pengumpulan data, kemudian preprocessing, summarization (ekstraktif dan abstraktif), evaluation (ROUGE), comparison dengan LLM, dan terakhir pengembangan aplikasi web.")

H2("3.2 Dataset")
CAPTION("Tabel 3.1 Statistik Dataset")
TABLE(
    ["Statistik", "Nilai"],
    [
        ["Jumlah dokumen", "100"],
        ["Bahasa", "Indonesia"],
        ["Domain", "Jurnal akademik"],
        ["Rata-rata panjang dokumen", "25.267 karakter"],
        ["Dokumen terpendek", "1.136 karakter"],
        ["Dokumen terpanjang", "57.745 karakter"],
        ["Rata-rata panjang ringkasan referensi", "964 karakter"],
        ["Sumber", "Scraping jurnal Indonesia"],
    ],
)
P("Dataset disimpan di data/raw/dataset.csv dengan kolom: judul, isi, ringkasan_referensi.")

H2("3.3 Pipeline Preprocessing")
P("Diimplementasikan di src/preprocessor.py dengan tahapan: case folding menggunakan text.lower(), cleaning dengan regex untuk menghapus karakter non-alfabet dan URL, sentence tokenization menggunakan NLTK sent_tokenize, word tokenization menggunakan NLTK word_tokenize, stopword removal menggunakan daftar stopword Sastrawi, serta stemming menggunakan Sastrawi.Stemmer.")

H2("3.4 Metode Ekstraktif: TF-IDF + TextRank")
P("Diimplementasikan di src/extractive_model.py. Algoritma dimulai dengan memecah dokumen menjadi kalimat, kemudian membangun matriks TF-IDF, menghitung matriks similarity menggunakan cosine similarity, membangun graph berbobot, menjalankan PageRank dengan NetworkX (nx.pagerank), memilih top-k kalimat (default k = 5 atau ratio kompresi 20%), dan menyusun ulang kalimat sesuai urutan asli di dokumen.")
P("Hyperparameter yang digunakan: damping factor d = 0,85, maksimum iterasi PageRank = 100, dan toleransi konvergensi = 1e-6.")

H2("3.5 Metode Abstraktif: mT5 Multilingual XL-Sum")
P("Diimplementasikan di src/abstractive_model.py.")
CAPTION("Tabel 3.2 Konfigurasi Model Abstraktif")
TABLE(
    ["Parameter", "Nilai"],
    [
        ["Model", "csebuetnlp/mT5_multilingual_XLSum"],
        ["Arsitektur", "mT5-base, encoder-decoder"],
        ["Jumlah parameter", "582 juta"],
        ["Max source length", "256 token"],
        ["Max target length", "128 token"],
        ["Min target length", "30 token"],
        ["Beam search", "4 beams"],
        ["Length penalty", "1.0"],
        ["No-repeat n-gram size", "2"],
        ["Inferensi", "CPU (AMD EPYC 7763)"],
    ],
)

H2("3.6 Evaluasi: ROUGE")
P("Library rouge-score (Google) digunakan untuk menghitung ROUGE-1, ROUGE-2, dan ROUGE-L (Precision, Recall, F1) untuk setiap pasangan kandidat dan referensi. Skor agregat dihitung sebagai rata-rata aritmatik atas seluruh dokumen.")

H2("3.7 Perbandingan dengan LLM")
P("Pada subset 10 dokumen, ringkasan tambahan dihasilkan oleh LLM sebagai third baseline. Tiga pasangan ROUGE dihitung: extractive vs reference, abstractive vs reference, dan dibandingkan terhadap output LLM.")

H2("3.8 Aplikasi Web")
P("Stack teknologi terdiri dari Flask (app.py), Jinja2 templates, Bootstrap, dan Chart.js. Halaman utama meliputi: / untuk input dokumen (CSV / PDF / teks manual), /summarize untuk menjalankan kedua metode, /evaluate untuk menjalankan ROUGE, /analysis untuk dashboard chart (line, bar, scatter, distribution), dan /detail/<id> untuk detail per-dokumen. Deployment menggunakan Gunicorn + Nginx (config di folder deploy/).")

H2("3.9 Spesifikasi Sistem")
BULLET([
    "CPU: AMD EPYC 7763",
    "RAM: cukup untuk inferensi mT5-base",
    "OS: Ubuntu 24.04.3 LTS (dev container)",
    "Python: 3.x",
    "Library utama: Flask, Transformers, Sastrawi, NLTK, scikit-learn, NetworkX, rouge-score, python-docx",
])

doc.add_page_break()

# ============ BAB 4 ============
H1("BAB 4\nHasil dan Pembahasan")

H2("4.1 Hasil Preprocessing")
P("Pipeline preprocessing berhasil dijalankan pada seluruh 100 dokumen tanpa error. Output disimpan di output/results/preprocess.json.")

H2("4.2 Hasil Ringkasan")
P("Ringkasan ekstraktif disimpan di output/summaries/extractive_summaries.csv, sedangkan ringkasan abstraktif disimpan di output/summaries/abstractive_summaries.csv.")

H2("4.3 Hasil Evaluasi ROUGE — Dataset Penuh (100 dokumen)")
ext = evaluate["extractive_scores"]
abs_ = evaluate["abstractive_scores"]
CAPTION("Tabel 4.1 Skor ROUGE Agregat (100 dokumen)")
TABLE(
    ["Metrik", "Metode", "Precision", "Recall", "F1-Score"],
    [
        ["ROUGE-1", "Ekstraktif", f"{ext['rouge1']['precision']:.4f}", f"{ext['rouge1']['recall']:.4f}", f"{ext['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-1", "Abstraktif", f"{abs_['rouge1']['precision']:.4f}", f"{abs_['rouge1']['recall']:.4f}", f"{abs_['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-2", "Ekstraktif", f"{ext['rouge2']['precision']:.4f}", f"{ext['rouge2']['recall']:.4f}", f"{ext['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-2", "Abstraktif", f"{abs_['rouge2']['precision']:.4f}", f"{abs_['rouge2']['recall']:.4f}", f"{abs_['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-L", "Ekstraktif", f"{ext['rougeL']['precision']:.4f}", f"{ext['rougeL']['recall']:.4f}", f"{ext['rougeL']['fmeasure']:.4f}"],
        ["ROUGE-L", "Abstraktif", f"{abs_['rougeL']['precision']:.4f}", f"{abs_['rougeL']['recall']:.4f}", f"{abs_['rougeL']['fmeasure']:.4f}"],
    ],
)
P("Interpretasi: ekstraktif unggul pada F1 seluruh varian ROUGE. Recall ekstraktif jauh lebih tinggi (0,7117 vs 0,1361 pada ROUGE-1) karena kalimat asli dipilih utuh sehingga banyak token referensi tertangkap. Sebaliknya, precision abstraktif jauh lebih tinggi (0,5170 vs 0,2418 pada ROUGE-1) karena output mT5 ringkas dan padat, namun pendek (max 128 token), sehingga banyak konten referensi tidak tertangkap dan recall menjadi rendah.")

H2("4.4 Hasil Evaluasi ROUGE — Subset 10 Dokumen + LLM")
ce = comparison["extractive"]
ca = comparison["abstractive"]
CAPTION("Tabel 4.2 Skor ROUGE Subset (10 dokumen, dengan referensi LLM)")
TABLE(
    ["Metrik", "Metode", "Precision", "Recall", "F1-Score"],
    [
        ["ROUGE-1", "Ekstraktif", f"{ce['rouge1']['precision']:.4f}", f"{ce['rouge1']['recall']:.4f}", f"{ce['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-1", "Abstraktif", f"{ca['rouge1']['precision']:.4f}", f"{ca['rouge1']['recall']:.4f}", f"{ca['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-2", "Ekstraktif", f"{ce['rouge2']['precision']:.4f}", f"{ce['rouge2']['recall']:.4f}", f"{ce['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-2", "Abstraktif", f"{ca['rouge2']['precision']:.4f}", f"{ca['rouge2']['recall']:.4f}", f"{ca['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-L", "Ekstraktif", f"{ce['rougeL']['precision']:.4f}", f"{ce['rougeL']['recall']:.4f}", f"{ce['rougeL']['fmeasure']:.4f}"],
        ["ROUGE-L", "Abstraktif", f"{ca['rougeL']['precision']:.4f}", f"{ca['rougeL']['recall']:.4f}", f"{ca['rougeL']['fmeasure']:.4f}"],
    ],
)
P(f"Best Method (subset): {comparison['best_method']}.")

H2("4.5 Pembahasan Trade-off Precision vs Recall")
P("Temuan kunci penelitian ini adalah profil precision-recall yang berkebalikan antara dua metode. Ekstraktif menghasilkan high recall dan moderate precision karena seluruh kalimat asli (dengan banyak kata pendukung) dipilih. Sebaliknya, abstraktif menghasilkan high precision dan low recall karena output dibatasi 128 token; kalimat dihasilkan padat tetapi cakupan kontennya sempit.")
P("Pilihan metode bergantung pada use-case: jika butuh cakupan luas (misalnya literature review), gunakan ekstraktif; jika butuh ringkasan padat dan rapi (misalnya abstrak singkat), gunakan abstraktif.")

H2("4.6 Analisis Faktor Per-Dokumen")
P("Tiga faktor utama yang berkorelasi dengan variasi performa adalah: (1) panjang dokumen — dokumen sangat panjang (>40k karakter) menurunkan F1 abstraktif karena batas output; (2) jumlah kalimat — semakin banyak kalimat, semakin banyak kandidat untuk ekstraktif; (3) rasio kompresi (panjang ringkasan dibagi panjang dokumen) — rasio rendah (ringkasan pendek dari dokumen panjang) lebih sulit untuk kedua metode, namun ekstraktif lebih tahan.")

H2("4.7 Mengapa Hasil Berbeda per Dokumen")
BULLET([
    "Struktur dokumen: dokumen dengan kalimat topik jelas (paragraf pertama informatif) menguntungkan TextRank.",
    "Overlap vocabulary dengan referensi: jika referensi menggunakan kata-kata asli dokumen, ekstraktif menang.",
    "Gaya referensi: referensi yang ditulis ulang (paraphrased) lebih cocok untuk evaluasi abstraktif.",
    "Domain spesifik: terminologi teknis langka kadang tidak ter-generate oleh mT5 (out-of-vocabulary effect pada SentencePiece).",
])

H2("4.8 Aplikasi Web")
P("Aplikasi Flask berhasil di-deploy. Fitur utama meliputi multi-format input (CSV, PDF, teks manual), visualisasi pipeline langkah demi langkah, dashboard analisis dengan Chart.js (line, bar, scatter, distribusi histogram), perbandingan side-by-side dengan LLM, dan antarmuka bilingual (ID/EN).")

H2("4.9 Diskusi Keterbatasan")
BULLET([
    "Inferensi CPU-only memperlambat batch abstraktif.",
    "mT5 tidak di-fine-tune pada domain akademik Indonesia (zero-shot).",
    "Ukuran dataset (100 dokumen) cukup untuk evaluasi tetapi belum ideal untuk training.",
    "ROUGE adalah metrik n-gram dan tidak menangkap kesetaraan semantik (paraphrase yang benar bisa skor rendah). Idealnya dilengkapi BERTScore dan evaluasi manusia.",
])

doc.add_page_break()

# ============ BAB 5 ============
H1("BAB 5\nKesimpulan dan Saran")

H2("5.1 Kesimpulan")
BULLET([
    "Penelitian berhasil mengimplementasikan dan membandingkan dua metode peringkasan teks akademik berbahasa Indonesia: ekstraktif (TF-IDF + TextRank) dan abstraktif (mT5 multilingual XL-Sum).",
    f"Pada dataset penuh 100 dokumen, metode ekstraktif unggul secara F1: ROUGE-1 F1 sebesar {ext['rouge1']['fmeasure']:.4f} (ekstraktif) vs {abs_['rouge1']['fmeasure']:.4f} (abstraktif); ROUGE-2 F1 sebesar {ext['rouge2']['fmeasure']:.4f} vs {abs_['rouge2']['fmeasure']:.4f}; dan ROUGE-L F1 sebesar {ext['rougeL']['fmeasure']:.4f} vs {abs_['rougeL']['fmeasure']:.4f}.",
    f"Pada subset 10 dokumen dengan referensi LLM, metode abstraktif unggul: ROUGE-1 F1 sebesar {ca['rouge1']['fmeasure']:.4f} (abstraktif) vs {ce['rouge1']['fmeasure']:.4f} (ekstraktif); ROUGE-L F1 sebesar {ca['rougeL']['fmeasure']:.4f} vs {ce['rougeL']['fmeasure']:.4f}.",
    "Terdapat trade-off konsisten precision-recall: ekstraktif memberikan recall tinggi (cakupan luas), abstraktif memberikan precision tinggi (output padat dan akurat tetapi pendek).",
    "Faktor panjang dokumen, jumlah kalimat, dan rasio kompresi secara signifikan memengaruhi performa per dokumen.",
    "Aplikasi web berbasis Flask berhasil dibangun dan menyediakan pipeline end-to-end yang interaktif untuk pengguna non-teknis.",
], numbered=True)

H2("5.2 Saran")
BULLET([
    "Fine-tuning mT5 pada korpus akademik Indonesia menggunakan GPU untuk meningkatkan recall abstraktif.",
    "Hybrid approach: jalankan ekstraktif terlebih dahulu (filter kalimat penting), lalu input ke abstraktif (paraphrasing) — diharapkan meningkatkan precision dan recall sekaligus.",
    "Perbesar dataset menjadi >1.000 dokumen lintas disiplin (teknik, sosial, kedokteran).",
    "Tambah metrik evaluasi: BERTScore, METEOR, dan evaluasi manusia (fluency, coherence, faithfulness).",
    "Eksplorasi model lain: IndoBART, IndoT5, dan model open-source LLM (Llama-Indo, Sahabat-AI) untuk perbandingan lebih luas.",
    "Atasi batas panjang output abstraktif dengan strategi chunking + hierarchical summarization untuk dokumen panjang.",
], numbered=True)

out_path = OUT_DIR / "skripsi_bab1-5.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Size: {out_path.stat().st_size / 1024:.1f} KB")

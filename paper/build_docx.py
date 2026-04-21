"""Build IEEE-style Word document from research findings using python-docx."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "paper"
OUT_DIR.mkdir(exist_ok=True)

# Load data
evaluate = json.loads((ROOT / "output/results/evaluate.json").read_text())
comparison = json.loads((ROOT / "output/results/comparison_summary.json").read_text())

doc = Document()

# Page setup (IEEE-like: Letter, 0.75" margins, 2 cols later if needed)
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)


def add_heading(text, level=1, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_para(text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
    for r_idx, row in enumerate(rows, start=1):
        cells = tbl.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
    return tbl


# ===== Title =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run(
    "Comparative Analysis of Extractive and Abstractive Text Summarization Methods for Indonesian Academic Documents"
)
trun.bold = True
trun.font.size = Pt(18)
trun.font.name = "Times New Roman"

# Authors
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
arun = auth.add_run("Author Name, Supervisor Name\nDepartment of Informatics, University Name\nCity, Indonesia\nemail@university.ac.id")
arun.font.size = Pt(11)
arun.font.name = "Times New Roman"

doc.add_paragraph()

# ===== Abstract =====
add_heading("Abstract", level=1, size=11, bold=True)
add_para(
    "This paper presents a comparative study of extractive and abstractive text summarization "
    "approaches applied to Indonesian academic documents. The extractive method employs TF-IDF "
    "weighting combined with TextRank (PageRank-based graph ranking), while the abstractive "
    "method utilizes a pre-trained multilingual mT5 model (mT5_multilingual_XLSum) fine-tuned "
    "for summarization tasks. We evaluate both methods on a dataset of 100 Indonesian academic "
    "documents using ROUGE-1, ROUGE-2, and ROUGE-L metrics. Experimental results on the full "
    "dataset show that the extractive approach achieves a higher average ROUGE-1 F1 score "
    f"({evaluate['extractive_scores']['rouge1']['fmeasure']:.4f}) compared to the abstractive "
    f"approach ({evaluate['abstractive_scores']['rouge1']['fmeasure']:.4f}). However, on a "
    "10-document LLM-referenced subset, the abstractive method surpasses the extractive method "
    f"(F1 = {comparison['abstractive']['rouge1']['fmeasure']:.4f} vs "
    f"{comparison['extractive']['rouge1']['fmeasure']:.4f}). Per-document analysis reveals that "
    "document length, sentence count, and compression ratio significantly influence "
    "summarization quality. A web-based interactive tool is developed to facilitate the "
    "complete NLP pipeline from preprocessing through evaluation."
)
add_para("Keywords—text summarization, extractive, abstractive, TF-IDF, TextRank, mT5, ROUGE, Indonesian NLP.", italic=True)

# ===== I. Introduction =====
add_heading("I. INTRODUCTION", size=12, bold=True)
add_para(
    "Text summarization is a fundamental task in Natural Language Processing (NLP) that aims to "
    "condense a document into a shorter version while preserving the most important information. "
    "With the exponential growth of academic publications, automatic text summarization has "
    "become increasingly valuable for researchers to quickly grasp the content of scholarly "
    "articles."
)
add_para(
    "Two primary approaches exist for automatic text summarization: extractive and abstractive. "
    "Extractive summarization selects and concatenates the most important sentences from the "
    "original document, preserving the original wording. Abstractive summarization, in contrast, "
    "generates new sentences that may not appear in the source document, aiming to capture the "
    "essence of the text in a more human-like manner."
)
add_para(
    "While significant progress has been made in text summarization for English and other "
    "high-resource languages, research on Indonesian text summarization, particularly for "
    "academic documents, remains relatively limited. Indonesian presents unique challenges "
    "including its agglutinative morphology, the prevalence of affixed words, and the limited "
    "availability of pre-trained language models specifically optimized for Indonesian."
)
add_para("This paper makes the following contributions:")
for c in [
    "1) A comparative analysis of extractive (TF-IDF + TextRank) and abstractive (mT5) summarization methods on 100 Indonesian academic documents.",
    "2) A per-document factor analysis examining how document characteristics influence summarization quality.",
    "3) Integration of LLM-generated summaries as a third comparison baseline with full ROUGE evaluation.",
    "4) An interactive web-based tool for the complete summarization and evaluation pipeline.",
]:
    p = doc.add_paragraph(c)
    p.paragraph_format.left_indent = Inches(0.25)
    for r in p.runs:
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

# ===== II. Related Work =====
add_heading("II. RELATED WORK", size=12, bold=True)
add_heading("A. Extractive Summarization", size=11)
add_para(
    "Extractive methods identify the most salient sentences in a document. Early approaches "
    "relied on term frequency and positional features [1]. Mihalcea and Tarau [2] introduced "
    "TextRank, a graph-based ranking algorithm inspired by PageRank, which constructs a graph "
    "of sentences with edges weighted by similarity. TF-IDF [3] has been widely used to "
    "represent sentence importance in extractive systems."
)
add_heading("B. Abstractive Summarization", size=11)
add_para(
    "Abstractive methods generate novel text using sequence-to-sequence models. The Transformer "
    "architecture [4] enabled significant advances, with models such as BART [5], T5 [6], and "
    "their multilingual variants (mBART, mT5) supporting cross-lingual summarization. Hasan et "
    "al. [7] introduced XL-Sum and fine-tuned mT5 for multilingual summarization across 44 "
    "languages including Indonesian."
)
add_heading("C. Indonesian Text Summarization", size=11)
add_para(
    "Research on Indonesian text summarization has explored both extractive [8] and abstractive "
    "[9] approaches. However, comparative studies examining both methods on academic documents "
    "with standardized evaluation metrics remain scarce."
)
add_heading("D. Evaluation Metrics", size=11)
add_para(
    "ROUGE [10] is the standard metric for summarization evaluation, measuring n-gram overlap "
    "between system-generated and reference summaries. ROUGE-1, ROUGE-2, and ROUGE-L capture "
    "unigram overlap, bigram overlap, and longest common subsequence, respectively."
)

# ===== III. Methodology =====
add_heading("III. METHODOLOGY", size=12, bold=True)
add_heading("A. Dataset", size=11)
add_para(
    "The dataset consists of 100 Indonesian academic documents collected from Indonesian "
    "academic journals. Each document is paired with a human-written reference summary. Table I "
    "summarizes the dataset statistics."
)
add_para("TABLE I. DATASET STATISTICS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_table(
    ["Statistic", "Value"],
    [
        ["Number of documents", "100"],
        ["Average document length", "25,267 characters"],
        ["Minimum document length", "1,136 characters"],
        ["Maximum document length", "57,745 characters"],
        ["Average summary length", "964 characters"],
        ["Language", "Indonesian"],
    ],
)

add_heading("B. Preprocessing Pipeline", size=11)
add_para("The preprocessing pipeline consists of six sequential stages: (1) Case Folding, (2) Cleaning of non-alphabetic characters and URLs, (3) Sentence Tokenization, (4) Word Tokenization, (5) Stopword Removal using a curated Indonesian stopword list, and (6) Stemming using the Sastrawi stemmer for Indonesian morphological affixation.")

add_heading("C. Extractive Summarization: TF-IDF + TextRank", size=11)
add_para(
    "The extractive method combines TF-IDF weighting with TextRank graph-based ranking. "
    "Each sentence is represented as a TF-IDF vector. Cosine similarity is computed between all "
    "sentence pairs to construct a weighted adjacency matrix. The PageRank algorithm is applied "
    "to rank sentences by centrality. The top-ranked sentences are selected and ordered by their "
    "original position in the document."
)

add_heading("D. Abstractive Summarization: mT5", size=11)
add_para(
    "The abstractive method uses the pre-trained csebuetnlp/mT5_multilingual_XLSum model "
    "(mT5-base, 582M parameters), pre-trained on mC4 and fine-tuned on the XL-Sum dataset. "
    "Configuration: maximum source length of 256 tokens, maximum target length of 128 tokens, "
    "minimum target length of 30 tokens, beam search with 4 beams."
)

add_heading("E. Evaluation Metrics", size=11)
add_para(
    "We evaluate using three ROUGE metrics: ROUGE-1 (unigram overlap), ROUGE-2 (bigram overlap), "
    "and ROUGE-L (Longest Common Subsequence). Each metric reports Precision, Recall, and F1-score."
)

# ===== IV. Results =====
add_heading("IV. RESULTS AND DISCUSSION", size=12, bold=True)

add_heading("A. Aggregate ROUGE Scores (100 documents)", size=11)
add_para("Table II presents the aggregate ROUGE scores on the full dataset.")
add_para("TABLE II. AGGREGATE ROUGE SCORES (FULL DATASET)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ext = evaluate["extractive_scores"]
abs_ = evaluate["abstractive_scores"]
add_table(
    ["Metric", "Method", "Precision", "Recall", "F1-Score"],
    [
        ["ROUGE-1", "Extractive", f"{ext['rouge1']['precision']:.4f}", f"{ext['rouge1']['recall']:.4f}", f"{ext['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-1", "Abstractive", f"{abs_['rouge1']['precision']:.4f}", f"{abs_['rouge1']['recall']:.4f}", f"{abs_['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-2", "Extractive", f"{ext['rouge2']['precision']:.4f}", f"{ext['rouge2']['recall']:.4f}", f"{ext['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-2", "Abstractive", f"{abs_['rouge2']['precision']:.4f}", f"{abs_['rouge2']['recall']:.4f}", f"{abs_['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-L", "Extractive", f"{ext['rougeL']['precision']:.4f}", f"{ext['rougeL']['recall']:.4f}", f"{ext['rougeL']['fmeasure']:.4f}"],
        ["ROUGE-L", "Abstractive", f"{abs_['rougeL']['precision']:.4f}", f"{abs_['rougeL']['recall']:.4f}", f"{abs_['rougeL']['fmeasure']:.4f}"],
    ],
)
add_para(
    f"On the full dataset of 100 documents, the extractive method achieves a higher average "
    f"ROUGE-1 F1 score ({ext['rouge1']['fmeasure']:.4f}) compared to the abstractive method "
    f"({abs_['rouge1']['fmeasure']:.4f})."
)

add_heading("B. LLM-Referenced Subset (10 documents)", size=11)
add_para("Table III presents results on a 10-document subset where summaries are also generated by a Large Language Model (LLM) for three-way comparison.")
add_para("TABLE III. ROUGE SCORES ON LLM-REFERENCED SUBSET", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
ce = comparison["extractive"]
ca = comparison["abstractive"]
add_table(
    ["Metric", "Method", "Precision", "Recall", "F1-Score"],
    [
        ["ROUGE-1", "Extractive", f"{ce['rouge1']['precision']:.4f}", f"{ce['rouge1']['recall']:.4f}", f"{ce['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-1", "Abstractive", f"{ca['rouge1']['precision']:.4f}", f"{ca['rouge1']['recall']:.4f}", f"{ca['rouge1']['fmeasure']:.4f}"],
        ["ROUGE-2", "Extractive", f"{ce['rouge2']['precision']:.4f}", f"{ce['rouge2']['recall']:.4f}", f"{ce['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-2", "Abstractive", f"{ca['rouge2']['precision']:.4f}", f"{ca['rouge2']['recall']:.4f}", f"{ca['rouge2']['fmeasure']:.4f}"],
        ["ROUGE-L", "Extractive", f"{ce['rougeL']['precision']:.4f}", f"{ce['rougeL']['recall']:.4f}", f"{ce['rougeL']['fmeasure']:.4f}"],
        ["ROUGE-L", "Abstractive", f"{ca['rougeL']['precision']:.4f}", f"{ca['rougeL']['recall']:.4f}", f"{ca['rougeL']['fmeasure']:.4f}"],
    ],
)
add_para(f"Best Method on this subset: {comparison['best_method']}.")

add_heading("C. Precision vs. Recall Trade-off", size=11)
add_para(
    "A notable finding is the contrasting precision-recall profiles of the two methods. The "
    "extractive method achieves moderate precision and higher recall, indicating that selected "
    "sentences capture a broader range of reference content but include some irrelevant material. "
    "The abstractive method achieves high precision but lower recall, suggesting that generated "
    "summaries are concise and relevant but miss significant portions of the reference content. "
    "This trade-off is largely attributed to the fixed maximum target length (128 tokens) "
    "imposed on the abstractive model."
)

add_heading("D. Per-Document Factor Analysis", size=11)
add_para(
    "Analysis of per-document ROUGE scores reveals substantial variation across documents. Three "
    "key factors influence performance: (1) Document Length—longer documents reduce abstractive "
    "performance due to fixed output length; (2) Sentence Count—more sentences provide more "
    "candidates for extractive selection; (3) Compression Ratio—higher ratios favor extractive "
    "methods, which can directly select key sentences."
)

add_heading("E. Why Results Differ Per Document", size=11)
add_para(
    "Several factors contribute to per-document variation: document structure (well-structured "
    "documents with clear topic sentences favor extractive); vocabulary overlap between source "
    "and reference summary; reference summary style (extractive style favors extractive method); "
    "and domain specificity (technical vocabulary may challenge mT5)."
)

# ===== V. Web Application =====
add_heading("V. WEB APPLICATION", size=12, bold=True)
add_para(
    "An interactive web application was developed using Flask to demonstrate the complete "
    "pipeline. Key features include multi-format input (CSV, PDF, manual text), step-by-step "
    "pipeline visualization, an analysis dashboard with Chart.js visualizations (line charts, "
    "bar charts, scatter plots, distribution histograms), LLM-comparison capability with "
    "side-by-side ROUGE evaluation, and a bilingual (Indonesian/English) interface."
)

# ===== VI. Conclusion =====
add_heading("VI. CONCLUSION", size=12, bold=True)
add_para(
    f"This study presented a comparative analysis of extractive and abstractive text "
    f"summarization methods for Indonesian academic documents. On the full 100-document dataset, "
    f"the extractive method (TF-IDF + TextRank) outperformed the abstractive method (mT5) with "
    f"ROUGE-1 F1 scores of {ext['rouge1']['fmeasure']:.4f} vs {abs_['rouge1']['fmeasure']:.4f}. "
    f"However, on a 10-document LLM-referenced subset, the abstractive method achieved superior "
    f"results (F1 = {ca['rouge1']['fmeasure']:.4f}). The abstractive method consistently yields "
    "higher precision but lower recall due to output length constraints, while document "
    "characteristics significantly affect per-document performance."
)
add_para(
    "Future work includes: (a) fine-tuning mT5 on Indonesian academic texts with GPU resources; "
    "(b) exploring hybrid extractive-then-abstractive methods; (c) expanding the dataset; and "
    "(d) incorporating BERTScore and human evaluation."
)

# ===== References =====
add_heading("REFERENCES", size=12, bold=True)
refs = [
    '[1]\tH. P. Luhn, "The automatic creation of literature abstracts," IBM J. Res. Dev., vol. 2, no. 2, pp. 159–165, 1958.',
    '[2]\tR. Mihalcea and P. Tarau, "TextRank: Bringing order into texts," in Proc. EMNLP, 2004, pp. 404–411.',
    '[3]\tG. Salton and C. Buckley, "Term-weighting approaches in automatic text retrieval," Inf. Process. Manag., vol. 24, no. 5, pp. 513–523, 1988.',
    '[4]\tA. Vaswani et al., "Attention is all you need," in Proc. NeurIPS, 2017, pp. 5998–6008.',
    '[5]\tM. Lewis et al., "BART: Denoising sequence-to-sequence pre-training for natural language generation, translation, and comprehension," in Proc. ACL, 2020, pp. 7871–7880.',
    '[6]\tC. Raffel et al., "Exploring the limits of transfer learning with a unified text-to-text transformer," JMLR, vol. 21, pp. 1–67, 2020.',
    '[7]\tT. Hasan et al., "XL-Sum: Large-scale multilingual abstractive summarization for 44 languages," in Findings of ACL, 2021, pp. 4693–4703.',
    '[8]\tF. Koto, J. H. Lau, and T. Baldwin, "IndoLEM and IndoBERT: A benchmark dataset and pre-trained language model for Indonesian NLP," in Proc. COLING, 2020, pp. 757–770.',
    '[9]\tB. Wilie et al., "IndoNLU: Benchmark and resources for evaluating Indonesian natural language understanding," in Proc. AACL-IJCNLP, 2020, pp. 843–857.',
    '[10]\tC.-Y. Lin, "ROUGE: A package for automatic evaluation of summaries," in Text Summarization Branches Out, 2004, pp. 74–81.',
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

out_path = OUT_DIR / "ieee_paper.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Size: {out_path.stat().st_size / 1024:.1f} KB")
